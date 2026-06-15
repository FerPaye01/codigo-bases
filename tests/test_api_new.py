import sys
import requests
import sqlite3
import os

sys.path.append(".")
import database

def run_test():
    print("--- 1. Testing /generate_docx endpoint with Bienes - Normal template ---")
    
    payload = {
        "nomenclatura": "LP-001-2026-OSINERGMIN",
        "numero_expediente": "20260000411",
        "objeto": "Adquisición de Servidores de Alta Disponibilidad para el Data Center de Osinergmin",
        "plazo": "60 días calendario",
        "valor_estimado": 450000.0,
        "moneda": "Soles (S/.)",
        "unidad_solicitante": "GSTI",
        "responsable": "lrodriguez@osinergmin.gob.pe",
        "fecha_convocatoria": "15/06/2026",
        "fecha_consultas": "19/06/2026",
        "fecha_absolucion": "21/06/2026",
        "fecha_presentacion": "27/06/2026",
        "fecha_buenapro": "30/06/2026",
        "fuente_financiamiento": "Recursos Ordinarios (RO)",
        "requisitos_calificacion": "Experiencia mínima de 3 años.",
        "factores_evaluacion": "Soporte técnico certificado.",
        "plantilla_usada": "Bienes - Normal",
        "sistema_contratacion": "Precios Unitarios"
    }
    
    try:
        resp = requests.post("http://127.0.0.1:8000/generate_docx", json=payload, timeout=10)
        if resp.status_code == 200:
            output_file = "output_bienes_normal.docx"
            with open(output_file, "wb") as f:
                f.write(resp.content)
            print(f"✓ Word document successfully generated and saved to {output_file}")
            
            # Use docx library to read the generated file and verify the title
            import docx
            doc = docx.Document(output_file)
            first_para = doc.paragraphs[0].text
            print(f"✓ Document Title Verification: {first_para.replace(chr(10), ' | ')}")
            
            # Verify if Sistema de Contratación is inside the table
            table = doc.tables[0]
            found = False
            for row in table.rows:
                if "Sistema de Contratación" in row.cells[0].text:
                    print(f"✓ Table check: found '{row.cells[0].text}' with value '{row.cells[1].text}'")
                    found = True
                    break
            if not found:
                print("✗ Error: 'Sistema de Contratación' not found in table!")
        else:
            print(f"✗ Request failed with status {resp.status_code}: {resp.text}")
    except Exception as e:
        print(f"✗ Connection error: {e}")
        
    print("\n--- 2. Checking SQLite auditoria table to verify RNF-1 traceability ---")
    conn = database.get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM auditoria ORDER BY fecha DESC LIMIT 5")
    rows = cursor.fetchall()
    conn.close()
    
    if rows:
        print(f"Found {len(rows)} audit entries (newest first):")
        for idx, row in enumerate(rows):
            print(f"[{idx+1}] Usuario: {row['usuario_id']} | Acción: {row['accion']} | Exp: {row['numero_expediente']} | Detalles: {row['detalles']} | Fecha: {row['fecha']}")
    else:
        print("✗ No audit entries found in DB!")

if __name__ == "__main__":
    run_test()
