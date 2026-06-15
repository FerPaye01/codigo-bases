import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_docx_template(filename, title_text, is_abreviado=False):
    doc = docx.Document()
    
    # Configurar estilos de fuente por defecto a Arial (muy usado en Osinergmin)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Título principal
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tipo_bases = "BASES ESTÁNDAR" if not is_abreviado else "BASES ESTÁNDAR SIMPLIFICADAS (ABREVIADAS)"
    title_run = title.add_run(f"OSINERGMIN\nGERENCIA DE ADMINISTRACIÓN Y FINANZAS\n\n{tipo_bases}\n{title_text.upper()}")
    title_run.bold = True
    title_run.font.size = Pt(13)
    title_run.font.color.rgb = RGBColor(12, 74, 110) # Azul Osinergmin (#0c4a6e)
    
    doc.add_paragraph("\n")
    
    # Sección 1: Datos Generales
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("CAPÍTULO I: DISPOSICIONES GENERALES Y DATOS DEL PROCESO")
    h1_run.bold = True
    h1_run.font.size = Pt(12)
    h1_run.font.color.rgb = RGBColor(12, 74, 110)
    
    table_data = doc.add_table(rows=7, cols=2)
    table_data.style = 'Table Grid'
    
    headers = [
        ("Nomenclatura Oficial del Proceso", "{{ nomenclatura }}"),
        ("Número de Expediente SIGED", "{{ numero_expediente }}"),
        ("Objeto de la Contratación", "{{ objeto }}"),
        ("Sistema de Contratación", "{{ sistema_contratacion }}"),
        ("Unidad Orgánica Solicitante", "{{ unidad_solicitante }}"),
        ("Responsable Técnico", "{{ responsable }}"),
        ("Fuente de Financiamiento", "{{ fuente_financiamiento }}")
    ]
    
    for idx, (label, placeholder) in enumerate(headers):
        row = table_data.rows[idx]
        # Celda de etiqueta
        cell_label = row.cells[0]
        cell_label.paragraphs[0].add_run(label).bold = True
        # Celda de valor (placeholder)
        cell_val = row.cells[1]
        cell_val.paragraphs[0].add_run(placeholder)
        
    doc.add_paragraph("\n")
    
    # Sección 2: Cronograma
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("CAPÍTULO II: CRONOGRAMA DEL PROCEDIMIENTO DE SELECCIÓN")
    h2_run.bold = True
    h2_run.font.size = Pt(12)
    h2_run.font.color.rgb = RGBColor(12, 74, 110)
    
    table_cron = doc.add_table(rows=6, cols=2)
    table_cron.style = 'Table Grid'
    
    # Cabecera de cronograma
    hdr_cells = table_cron.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run("Etapa del Proceso").bold = True
    hdr_cells[1].paragraphs[0].add_run("Fecha programada").bold = True
    
    cron_steps = [
        ("1. Convocatoria y publicación en el SEACE", "{{ fecha_convocatoria }}"),
        ("2. Formulación de consultas y observaciones", "{{ fecha_consultas }}"),
        ("3. Absolución de consultas e integración de bases", "{{ fecha_absolucion }}"),
        ("4. Presentación de ofertas técnicas y económicas", "{{ fecha_presentacion }}"),
        ("5. Otorgamiento de la Buena Pro", "{{ fecha_buenapro }}")
    ]
    
    for idx, (step, placeholder) in enumerate(cron_steps):
        row = table_cron.rows[idx + 1]
        row.cells[0].paragraphs[0].add_run(step)
        row.cells[1].paragraphs[0].add_run(placeholder)
        
    doc.add_paragraph("\n")
    
    # Sección 3: Datos Económicos
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("CAPÍTULO III: VALOR ESTIMADO Y CONDICIONES ECONÓMICAS")
    h3_run.bold = True
    h3_run.font.size = Pt(12)
    h3_run.font.color.rgb = RGBColor(12, 74, 110)
    
    p_valor = doc.add_paragraph()
    p_valor.add_run("El valor estimado de la presente contratación asciende a la suma de ").font.size = Pt(11)
    p_valor.add_run("{{ moneda }} {{ valor_estimado }}").bold = True
    
    doc.add_paragraph("\n")
    
    # Sección 4: Requerimientos Técnicos
    h4 = doc.add_paragraph()
    h4_run = h4.add_run("CAPÍTULO IV: REQUERIMIENTOS TÉCNICOS MÍNIMOS (PLAZOS Y REQUISITOS)")
    h4_run.bold = True
    h4_run.font.size = Pt(12)
    h4_run.font.color.rgb = RGBColor(12, 74, 110)
    
    p_plazo = doc.add_paragraph()
    p_plazo.add_run("Plazo de ejecución: ").bold = True
    p_plazo.add_run("{{ plazo }}")
    
    doc.add_paragraph("Requisitos de Calificación Mínimos:").bold = True
    p_reqs = doc.add_paragraph()
    p_reqs.add_run("{{ requisitos_calificacion }}")
    
    doc.add_paragraph("\n")
    
    # Sección 5: Factores de Evaluación
    h5 = doc.add_paragraph()
    h5_run = h5.add_run("CAPÍTULO V: FACTORES DE EVALUACIÓN")
    h5_run.bold = True
    h5_run.font.size = Pt(12)
    h5_run.font.color.rgb = RGBColor(12, 74, 110)
    
    p_factores = doc.add_paragraph()
    p_factores.add_run("{{ factores_evaluacion }}")
    
    doc.save(filename)
    print(f"Plantilla '{filename}' creada exitosamente.")

def create_all_templates():
    # 1. Base / Fallback template
    generate_docx_template("plantilla_base.docx", "Bases de Procedimiento de Selección (Genérica)")
    
    # 2. Bienes
    generate_docx_template("plantilla_bienes_normal.docx", "Adquisición de Bienes", is_abreviado=False)
    generate_docx_template("plantilla_bienes_abreviado.docx", "Adquisición de Bienes", is_abreviado=True)
    
    # 3. Servicios
    generate_docx_template("plantilla_servicios_normal.docx", "Contratación de Servicios en General", is_abreviado=False)
    generate_docx_template("plantilla_servicios_abreviado.docx", "Contratación de Servicios en General", is_abreviado=True)
    
    # 4. Consultorías
    generate_docx_template("plantilla_consultoria_normal.docx", "Contratación de Servicios de Consultoría en General", is_abreviado=False)
    generate_docx_template("plantilla_consultoria_abreviado.docx", "Contratación de Servicios de Consultoría en General", is_abreviado=True)

if __name__ == "__main__":
    create_all_templates()
